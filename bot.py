import asyncio
import os
import re

from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.types import Message, ReplyKeyboardRemove, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery, FSInputFile
from dotenv import load_dotenv
from aiogram.types import FSInputFile
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side

from aiogram.types import (
    Message,
    ReplyKeyboardRemove,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    CallbackQuery,
    FSInputFile
)

from keyboards import user_keyboard, admin_keyboard
from db import (
    get_staff_by_telegram_id,
    create_staff,
    create_event,
    get_all_staff,
    get_active_event,
    add_event_response,
    count_event_responses,
    close_event,
    get_event_participants,
    confirm_event_participants,
    get_confirmed_participants,
    confirm_shift_participants,
    remove_participant,
    find_participant_by_lastname,
    cancel_event,
    clear_old_events,
    get_registered_staff,
    find_staff_by_lastname,
    add_staff_to_blacklist,
    get_blacklist,
    is_staff_blacklisted,
    remove_staff_from_blacklist,
    get_event_by_date,
    create_event_shift,
    get_event_shifts,
    get_shift_by_date_and_start,
    add_shift_response,
    count_shift_responses,
    close_shift,
    get_all_events_with_shifts,
    get_shift_participants,
    get_events_by_date_with_shifts,
    find_shift_participant_by_lastname,
    remove_participant_from_shift,
    close_event_shifts,
    get_event_participants_by_event,
    remove_staff_from_shift_response,
    has_staff_response_on_event_date,
    get_active_events_for_users,
    get_staff_active_responses,
    get_shift_by_id,
    get_tomorrow_shift_participants_for_reminder,
    mark_shift_reminder_sent,
    get_active_events,
    get_active_shifts_for_admin,
    get_staff_active_responses_with_ids,
    archive_past_events
)

load_dotenv()

BOT_TOKEN = os.getenv("BOT_TOKEN")
ADMIN_ID = os.getenv("ADMIN_ID")

if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN не найден в .env")

if not ADMIN_ID:
    raise ValueError("ADMIN_ID не найден в .env")

ADMIN_ID = int(ADMIN_ID)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()


class Registration(StatesGroup):
    full_name = State()
    citizenship = State()
    phone = State()
    passport_photo = State()


class CreateEvent(StatesGroup):
    event_date = State()
    comment = State()

class AddShift(StatesGroup):
    event_date = State()
    shift_start = State()
    shift_end = State()
    staff_required = State()


class RemoveParticipant(StatesGroup):
    lastname = State()


class CancelEvent(StatesGroup):
    event_date = State()
    confirm = State()

class Broadcast(StatesGroup):
    text = State()

class AddBlacklist(StatesGroup):
    lastname = State()
    reason = State()

class RemoveBlacklist(StatesGroup):
    lastname = State()

class StaffListByDate(StatesGroup):
    event_date = State()

class SignUpShift(StatesGroup):
    shift_data = State()

class ConfirmShiftParticipants(StatesGroup):
    event_date = State()
    shift_start = State()

class RemoveParticipantFromShift(StatesGroup):
    event_date = State()
    shift_start = State()
    lastname = State()

class RefuseShift(StatesGroup):
    event_date = State()
    shift_start = State()


@dp.message(CommandStart())
async def start_handler(message: Message, state: FSMContext):
    telegram_id = message.from_user.id

    if telegram_id == ADMIN_ID:
        await state.clear()
        await message.answer(
            "Здравствуйте! Вы вошли как администратор.",
            reply_markup=admin_keyboard()
        )
        return

    staff = get_staff_by_telegram_id(telegram_id)

    if staff and is_staff_blacklisted(staff["staff_id"]):
        await message.answer(
            "❌ Вы не можете пользоваться ботом, так как находитесь в черном списке."
        )
        return

    if staff:
        await state.clear()
        await message.answer(
            "Вы уже зарегистрированы в системе.",
            reply_markup=user_keyboard()
        )
        return

    await message.answer(
        "Приветствую тебя в нашей команде кейтеринга by Demidov!\n\n"
        "Для начала работы необходимо указать данные о себе.\n\n"
        "Напишите полностью ФИО:",
        reply_markup=ReplyKeyboardRemove()
    )

    await state.set_state(Registration.full_name)


@dp.message(Registration.full_name)
async def process_full_name(message: Message, state: FSMContext):
    await state.update_data(full_name=message.text)
    await message.answer("Укажите гражданство:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(Registration.citizenship)


@dp.message(Registration.citizenship)
async def process_citizenship(message: Message, state: FSMContext):
    await state.update_data(citizenship=message.text)
    await message.answer("Укажите номер телефона:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(Registration.phone)


@dp.message(Registration.phone)
async def process_phone(message: Message, state: FSMContext):
    await state.update_data(phone=message.text)
    await message.answer("Пришлите фото паспорта:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(Registration.passport_photo)


@dp.message(Registration.passport_photo, F.photo)
async def process_passport_photo(message: Message, state: FSMContext):
    data = await state.get_data()
    passport_photo_id = message.photo[-1].file_id

    create_staff(
        full_name=data["full_name"],
        citizenship=data["citizenship"],
        phone=data["phone"],
        telegram_id=message.from_user.id,
        telegram_username=message.from_user.username,
        passport_photo_id=passport_photo_id
    )

    await state.clear()

    await message.answer(
        "Спасибо! Ваши данные сохранены.\n\n"
        "Теперь вы добавлены в общий список персонала и сможете получать рассылки о мероприятиях.",
        reply_markup=user_keyboard()
    )

@dp.message(F.text == "📋 Активные мероприятия")
async def active_events_user_handler(message: Message):
    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer(
            "Сначала необходимо пройти регистрацию через команду /start.",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    if is_staff_blacklisted(staff["staff_id"]):
        await message.answer(
            "❌ Вы не можете просматривать мероприятия, так как находитесь в черном списке."
        )
        return

    rows = get_active_events_for_users()

    if not rows:
        await message.answer("Сейчас нет активных мероприятий для записи.")
        return

    text = "📋 Активные мероприятия и смены:\n\n"

    current_event_id = None

    for row in rows:
        if row["event_id"] != current_event_id:
            current_event_id = row["event_id"]

            text += (
                f"📅 Дата: {row['event_date']}\n"
                f"Комментарий: {row['comment']}\n\n"
            )

        free_places = row["staff_required"] - row["response_count"]

        text += (
            f"🕒 Смена: {row['shift_start']}–{row['shift_end']}\n"
            f"Нужно сотрудников: {row['staff_required']}\n"
            f"Уже записались: {row['response_count']}\n"
            f"Свободных мест: {free_places}\n"
            f"Для записи укажите: {row['event_date']} {row['shift_start']}\n\n"
        )

    text += (
        "Чтобы записаться, нажмите кнопку «Записаться на смену» "
        "и отправьте дату с временем начала смены."
    )

    await message.answer(text)

@dp.message(F.text == "📌 Мои записи")
async def my_records_handler(message: Message):
    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer(
            "Сначала необходимо пройти регистрацию через команду /start.",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    if is_staff_blacklisted(staff["staff_id"]):
        await message.answer(
            "❌ Вы не можете просматривать записи, так как находитесь в черном списке."
        )
        return

    rows = get_staff_active_responses(staff["staff_id"])

    if not rows:
        await message.answer("У вас пока нет активных записей на мероприятия.")
        return

    text = "📌 Ваши активные записи:\n\n"

    for i, row in enumerate(rows, start=1):
        text += (
            f"{i}. 📅 Дата: {row['event_date']}\n"
            f"🕒 Смена: {row['shift_start']}–{row['shift_end']}\n"
            f"Статус: {row['response_status']}\n"
            f"Комментарий: {row['comment']}\n\n"
        )

    await message.answer(text)

@dp.message(F.text == "✅ Записаться на смену")
async def signup_shift_start(message: Message, state: FSMContext):
    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer(
            "Сначала необходимо пройти регистрацию через команду /start.",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    if is_staff_blacklisted(staff["staff_id"]):
        await message.answer(
            "❌ Вы не можете записываться на мероприятия, так как находитесь в черном списке."
        )
        return

    rows = get_active_events_for_users()

    if not rows:
        await message.answer("Сейчас нет активных мероприятий для записи.")
        return

    text = "Выберите смену для записи:\n\n"

    keyboard = []

    current_event_id = None

    for row in rows:
        if row["event_id"] != current_event_id:
            current_event_id = row["event_id"]
            text += (
                f"📅 Дата: {row['event_date']}\n"
                f"Комментарий: {row['comment']}\n\n"
            )

        free_places = row["staff_required"] - row["response_count"]

        text += (
            f"🕒 Смена: {row['shift_start']}–{row['shift_end']}\n"
            f"Свободных мест: {free_places}\n\n"
        )

        if free_places > 0:
            keyboard.append([
                InlineKeyboardButton(
                    text=f"Записаться {row['event_date']} {row['shift_start']}",
                    callback_data=f"signup_shift:{row['shift_id']}"
                )
            ])

    if not keyboard:
        await message.answer("Сейчас нет смен со свободными местами.")
        return

    await message.answer(
        text,
        reply_markup=InlineKeyboardMarkup(inline_keyboard=keyboard)
    )

@dp.callback_query(F.data.startswith("signup_shift:"))
async def signup_shift_callback(callback: CallbackQuery):
    shift_id = int(callback.data.split(":")[1])

    staff = get_staff_by_telegram_id(callback.from_user.id)

    if not staff:
        await callback.message.answer("Сначала необходимо пройти регистрацию через команду /start.")
        await callback.answer()
        return

    if is_staff_blacklisted(staff["staff_id"]):
        await callback.message.answer(
            "❌ Вы не можете записываться на мероприятия, так как находитесь в черном списке."
        )
        await callback.answer()
        return

    shift = get_shift_by_id(shift_id)

    if not shift:
        await callback.message.answer("Эта смена уже недоступна для записи.")
        await callback.answer()
        return

    if has_staff_response_on_event_date(shift["event_date"], staff["staff_id"]):
        await callback.message.answer(
            "Вы уже записаны на смену в этот день.\n\n"
            "В один день можно записаться только на одну смену."
        )
        await callback.answer()
        return

    current_count = count_shift_responses(shift["shift_id"])
    required_count = shift["staff_required"]

    if current_count >= required_count:
        close_shift(shift["shift_id"])
        await callback.message.answer("На эту смену набор уже закрыт.")
        await callback.answer()
        return

    added = add_shift_response(
        event_id=shift["event_id"],
        shift_id=shift["shift_id"],
        staff_id=staff["staff_id"]
    )

    if not added:
        await callback.message.answer("Вы уже записались на эту смену.")
        await callback.answer()
        return

    current_count = count_shift_responses(shift["shift_id"])

    await callback.message.answer(
        "✅ Вы записались на смену.\n\n"
        f"Дата мероприятия: {shift['event_date']}\n"
        f"Смена: {shift['shift_start']}–{shift['shift_end']}\n"
        f"Сейчас записалось: {current_count} из {required_count}"
    )

    await bot.send_message(
        ADMIN_ID,
        "Новый отклик на смену.\n\n"
        f"ФИО: {staff['full_name']}\n"
        f"Телефон: {staff['phone']}\n"
        f"Дата: {shift['event_date']}\n"
        f"Смена: {shift['shift_start']}–{shift['shift_end']}\n"
        f"Записалось: {current_count} из {required_count}"
    )

    if current_count >= required_count:
        close_shift(shift["shift_id"])

        staff_list = get_all_staff()

        for person in staff_list:
            try:
                await bot.send_message(
                    person["telegram_id"],
                    "🔒 Набор на мероприятие закрыт.\n\n"
                    f"Дата мероприятия: {shift['event_date']}\n"
                    f"Смена: {shift['shift_start']}–{shift['shift_end']}"
                )
            except Exception:
                pass

        await bot.send_message(
            ADMIN_ID,
            "🔒 Нужное количество персонала на смену набрано.\n\n"
            f"Дата: {shift['event_date']}\n"
            f"Смена: {shift['shift_start']}–{shift['shift_end']}\n\n"
            "Уведомление отправлено всем пользователям."
        )

    await callback.answer("Запись выполнена")

@dp.message(Registration.passport_photo)
async def passport_photo_error(message: Message):
    await message.answer("Пожалуйста, отправьте именно фото паспорта.")

@dp.message(F.text == "📅 Создать мероприятие")
async def create_event_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    await message.answer(
        "Введите дату мероприятия в формате ГГГГ-ММ-ДД.\n"
        "Например: 2026-05-20"
    )
    await state.set_state(CreateEvent.event_date)


@dp.message(CreateEvent.event_date)
async def process_event_date(message: Message, state: FSMContext):
    await state.update_data(event_date=message.text)

    await message.answer(
        "Введите комментарий к мероприятию.\n"
        "Например: Банкет на 150 гостей, площадка — Москва."
    )
    await state.set_state(CreateEvent.comment)


@dp.message(CreateEvent.comment)
async def process_event_comment(message: Message, state: FSMContext):
    data = await state.get_data()

    event_id = create_event(
        event_date=data["event_date"],
        comment=message.text
    )

    await state.clear()

    await message.answer(
        "✅ Мероприятие создано.\n\n"
        f"ID мероприятия: {event_id}\n"
        f"Дата: {data['event_date']}\n"
        f"Комментарий: {message.text}\n\n"
        "Теперь можно добавить смену через кнопку «Добавить смену»."
    )

@dp.message(F.text.lower().in_(["+", "да"]))
async def event_response_handler(message: Message):
    await message.answer(
        "Чтобы записаться на смену, отправьте дату мероприятия и время начала смены.\n\n"
        "Пример:\n"
        "2026-05-20 07:00"
    )


@dp.message(F.text == "👥 Список персонала мероприятия")
async def get_participants_handler(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    await message.answer(
        "Введите дату мероприятия, по которому нужно посмотреть персонал.\n"
        "Формат: ГГГГ-ММ-ДД\n"
        "Например: 2026-05-20"
    )

    await state.set_state(StaffListByDate.event_date)

@dp.message(StaffListByDate.event_date)
async def staff_list_by_date_handler(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    event_date = message.text.strip()

    rows = get_events_by_date_with_shifts(event_date)

    if not rows:
        await message.answer("Мероприятие на эту дату не найдено.")
        await state.clear()
        return

    text = f"👥 Персонал мероприятия на дату {event_date}:\n\n"

    for row in rows:
        if not row["shift_id"]:
            text += "Смены пока не добавлены.\n\n"
            continue

        participants = get_shift_participants(row["shift_id"])

        text += (
            f"🕒 Смена: {row['shift_start']}–{row['shift_end']}\n"
            f"Нужно сотрудников: {row['staff_required']}\n"
            f"Записалось: {len(participants)}\n\n"
        )

        if participants:
            for i, participant in enumerate(participants, start=1):
                text += (
                    f"{i}. {participant['full_name']}\n"
                    f"Телефон: {participant['phone']}\n"
                    f"Статус: {participant['response_status']}\n\n"
                )
        else:
            text += "Пока никто не записался.\n\n"

    await state.clear()

    await message.answer(text)

    file_name = create_staff_docx(event_date, rows)

    await message.answer_document(
        FSInputFile(file_name),
        caption="📄 Word-файл со списком персонала"
    )

    excel_file = create_staff_xlsx(event_date, rows)

    await message.answer_document(
        FSInputFile(excel_file),
        caption="📊 Excel-файл со списком персонала"
    )

@dp.message(F.text == "👤 Зарегистрированные пользователи")
async def registered_staff_handler(message: Message):
    if message.from_user.id != ADMIN_ID:
        return

    staff_list = get_registered_staff()

    if not staff_list:
        await message.answer("Зарегистрированных пользователей пока нет.")
        return

    text = "Зарегистрированные пользователи:\n\n"

    for i, staff in enumerate(staff_list, start=1):
        username = staff["telegram_username"] or "не указан"
        text += (
            f"{i}. {staff['full_name']}\n"
            f"Телефон: {staff['phone']}\n"
            f"Гражданство: {staff['citizenship']}\n"
            f"Telegram: @{username}\n\n"
        )

    await message.answer(text)


@dp.message(F.text == "🚫 Добавить в черный список")
async def add_blacklist_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    await message.answer("Введите фамилию пользователя, которого нужно добавить в черный список:")
    await state.set_state(AddBlacklist.lastname)


@dp.message(AddBlacklist.lastname)
async def add_blacklist_lastname(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    staff = find_staff_by_lastname(message.text.strip())

    if not staff:
        await message.answer("Пользователь с такой фамилией не найден. Попробуйте снова.")
        return

    await state.update_data(
    staff_id=staff["staff_id"],
    full_name=staff["full_name"],
    telegram_id=staff["telegram_id"]  # ← ВОТ ЭТО ДОБАВЛЯЕМ
    )

    await message.answer(
        f"Найден пользователь: {staff['full_name']}\n"
        f"Телефон: {staff['phone']}\n\n"
        "Укажите причину добавления в черный список:"
    )
    await state.set_state(AddBlacklist.reason)


@dp.message(AddBlacklist.reason)
async def add_blacklist_reason(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    data = await state.get_data()

    add_staff_to_blacklist(
        staff_id=data["staff_id"],
        reason=message.text
    )

    try:
        await bot.send_message(
            data["telegram_id"],
            "❌ Вы добавлены в черный список.\n\n"
            "Вы больше не сможете записываться на мероприятия."
        )
    except Exception:
        pass

    await state.clear()

    await message.answer(
        f"Пользователь {data['full_name']} добавлен в черный список."
    )


@dp.message(F.text == "📛 Черный список")
async def blacklist_handler(message: Message):
    if message.from_user.id != ADMIN_ID:
        return

    blacklist = get_blacklist()

    if not blacklist:
        await message.answer("Черный список пуст.")
        return

    text = "Черный список:\n\n"

    for i, item in enumerate(blacklist, start=1):
        username = item["telegram_username"] or "не указан"
        text += (
            f"{i}. {item['full_name']}\n"
            f"Телефон: {item['phone']}\n"
            f"Telegram: @{username}\n"
            f"Причина: {item['reason']}\n\n"
        )

    await message.answer(text)


@dp.message(F.text == "🗑 Удалить участника")
async def remove_participant_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    await message.answer(
        "Введите дату мероприятия, из которого нужно удалить сотрудника.\n"
        "Формат: ГГГГ-ММ-ДД\n"
        "Например: 2026-05-20"
    )
    await state.set_state(RemoveParticipantFromShift.event_date)


@dp.message(RemoveParticipantFromShift.event_date)
async def remove_participant_event_date(message: Message, state: FSMContext):
    await state.update_data(event_date=message.text.strip())

    await message.answer(
        "Введите время начала смены.\n"
        "Формат: ЧЧ:ММ\n"
        "Например: 07:00"
    )
    await state.set_state(RemoveParticipantFromShift.shift_start)


@dp.message(RemoveParticipantFromShift.shift_start)
async def remove_participant_shift_start(message: Message, state: FSMContext):
    data = await state.get_data()

    shift = get_shift_by_date_and_start(
        event_date=data["event_date"],
        shift_start=message.text.strip()
    )

    if not shift:
        await message.answer("Смена на указанную дату и время не найдена.")
        await state.clear()
        return

    await state.update_data(
        shift_id=shift["shift_id"],
        shift_start=message.text.strip(),
        shift_end=shift["shift_end"]
    )

    participants = get_shift_participants(shift["shift_id"])

    if not participants:
        await message.answer("На этой смене пока нет записанных сотрудников.")
        await state.clear()
        return

    text = "Введите фамилию сотрудника, которого нужно удалить.\n\nТекущий список смены:\n\n"

    for i, participant in enumerate(participants, start=1):
        text += f"{i}. {participant['full_name']} — {participant['phone']}\n"

    await message.answer(text)
    await state.set_state(RemoveParticipantFromShift.lastname)


@dp.message(RemoveParticipantFromShift.lastname)
async def remove_participant_lastname(message: Message, state: FSMContext):
    data = await state.get_data()

    participant = find_shift_participant_by_lastname(
        shift_id=data["shift_id"],
        lastname=message.text.strip()
    )

    if not participant:
        await message.answer("Сотрудник с такой фамилией не найден на данной смене.")
        return

    remove_participant_from_shift(
        shift_id=data["shift_id"],
        staff_id=participant["staff_id"]
    )

    try:
        await bot.send_message(
            participant["telegram_id"],
            "❌ Вас удалили из списка персонала мероприятия.\n\n"
            f"Дата: {data['event_date']}\n"
            f"Смена: {data['shift_start']}–{data['shift_end']}"
        )
    except Exception:
        pass

    await state.clear()

    await message.answer(
        f"Сотрудник {participant['full_name']} удалён из смены.\n\n"
        f"Дата: {data['event_date']}\n"
        f"Смена: {data['shift_start']}–{data['shift_end']}"
    )


@dp.message(RemoveParticipant.lastname)
async def remove_participant_by_lastname(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    active_event = get_active_event()

    if not active_event:
        await message.answer("Нет активного мероприятия.")
        await state.clear()
        return

    lastname = message.text.strip()

    participant = find_participant_by_lastname(
        event_id=active_event["event_id"],
        lastname=lastname
    )

    if not participant:
        await message.answer(
            "Участник с такой фамилией не найден в списке.\n"
            "Проверьте написание фамилии и попробуйте снова."
        )
        return

    remove_participant(active_event["event_id"], participant["staff_id"])

    try:
        await bot.send_message(
            participant["telegram_id"],
            "❌ Вас НЕ записали на мероприятие.\n\n"
            f"Дата: {active_event['event_date']}\n"
            f"Время сбора: {active_event['start_time']}"
        )
    except Exception:
        pass

    await state.clear()

    await message.answer(f"Участник {participant['full_name']} удалён из списка.")

@dp.message(F.text == "✅ Подтвердить список")
async def confirm_participants_start(message: Message):
    if message.from_user.id != ADMIN_ID:
        return

    shifts = get_active_shifts_for_admin()

    if not shifts:
        await message.answer("Нет активных смен для подтверждения.")
        return

    text = "Выберите смену, по которой нужно подтвердить список:\n\n"
    keyboard = []

    for shift in shifts:
        text += (
            f"📅 {shift['event_date']}\n"
            f"🕒 {shift['shift_start']}–{shift['shift_end']}\n"
            f"Комментарий: {shift['comment']}\n"
            f"Записано: {shift['response_count']} из {shift['staff_required']}\n\n"
        )

        keyboard.append([
            InlineKeyboardButton(
                text=f"{shift['event_date']} {shift['shift_start']}–{shift['shift_end']}",
                callback_data=f"confirm_shift:{shift['shift_id']}"
            )
        ])

    await message.answer(
        text,
        reply_markup=InlineKeyboardMarkup(inline_keyboard=keyboard)
    )

@dp.callback_query(F.data.startswith("confirm_shift:"))
async def confirm_shift_callback(callback: CallbackQuery):
    if callback.from_user.id != ADMIN_ID:
        await callback.answer()
        return

    shift_id = int(callback.data.split(":")[1])

    shift = get_shift_by_id(shift_id)

    if not shift:
        await callback.message.answer("Смена не найдена или уже закрыта.")
        await callback.answer()
        return

    participants = get_shift_participants(shift_id)

    if not participants:
        await callback.message.answer("На этой смене пока нет сотрудников для подтверждения.")
        await callback.answer()
        return

    confirm_shift_participants(shift_id)

    confirmed_participants = get_shift_participants(shift_id)

    for participant in confirmed_participants:
        try:
            await bot.send_message(
                participant["telegram_id"],
                "✅ Вас подтвердили на смену.\n\n"
                f"Дата: {shift['event_date']}\n"
                f"Смена: {shift['shift_start']}–{shift['shift_end']}\n\n"
                "Пожалуйста, не опаздывайте."
            )
        except Exception:
            pass

    text = "✅ Список смены подтверждён.\n\n"
    text += f"Дата: {shift['event_date']}\n"
    text += f"Смена: {shift['shift_start']}–{shift['shift_end']}\n\n"
    text += "Подтверждённые сотрудники:\n\n"

    for i, participant in enumerate(confirmed_participants, start=1):
        text += f"{i}. {participant['full_name']} — {participant['phone']}\n"

    await callback.message.answer(text)
    await callback.answer("Список подтверждён")

@dp.message(F.text == "❌ Отмена мероприятия")
async def cancel_event_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    await message.answer(
        "Введите дату мероприятия, которое нужно отменить.\n"
        "Формат: ГГГГ-ММ-ДД\n"
        "Например: 2026-05-20"
    )

    await state.set_state(CancelEvent.event_date)

@dp.message(CancelEvent.event_date)
async def cancel_event_date(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    event_date = message.text.strip()

    event = get_event_by_date(event_date)

    if not event:
        await message.answer("Активное мероприятие на эту дату не найдено.")
        await state.clear()
        return

    await state.update_data(
        event_id=event["event_id"],
        event_date=event["event_date"],
        comment=event["comment"]
    )

    await message.answer(
        "Вы действительно хотите отменить мероприятие?\n\n"
        f"Дата: {event['event_date']}\n"
        f"Комментарий: {event['comment']}\n\n"
        "Для подтверждения напишите: ПОДТВЕРДИТЬ\n"
        "Для отмены напишите: НЕТ"
    )

    await state.set_state(CancelEvent.confirm)

@dp.message(CancelEvent.confirm)
async def cancel_event_confirm(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    answer = message.text.strip().lower()

    if answer == "нет":
        await message.answer("Отмена мероприятия не выполнена.")
        await state.clear()
        return

    if answer != "подтвердить":
        await message.answer(
            "Команда не распознана.\n"
            "Напишите ПОДТВЕРДИТЬ для отмены мероприятия или НЕТ для выхода."
        )
        return

    data = await state.get_data()

    cancel_event(data["event_id"])
    close_event_shifts(data["event_id"])

    participants = get_event_participants_by_event(data["event_id"])

    for participant in participants:
        try:
            await bot.send_message(
                participant["telegram_id"],
                "⚠️ Мероприятие отменено.\n\n"
                f"Дата: {data['event_date']}\n"
                f"Комментарий: {data['comment']}\n\n"
                "Просим не выходить на данное мероприятие."
            )
        except Exception:
            pass

    await state.clear()

    await message.answer(
        "✅ Мероприятие отменено.\n"
        f"Уведомление отправлено участникам: {len(participants)}"
    )


@dp.message(F.text == "📢 Массовый комментарий")
async def broadcast_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    await message.answer("Введите текст для массовой рассылки:")
    await state.set_state(Broadcast.text)


@dp.message(Broadcast.text)
async def broadcast_send(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    text = message.text
    staff_list = get_all_staff()
    count = 0

    for person in staff_list:
        try:
            await bot.send_message(person["telegram_id"], text)
            count += 1
        except Exception:
            pass

    await state.clear()

    await message.answer(f"Сообщение отправлено пользователям: {count}")


@dp.message(F.text == "📞 Связь с менеджером")
async def manager_contact(message: Message):
    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer(
            "Сначала необходимо пройти регистрацию через команду /start.",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    await message.answer(
        "📞 Контакты менеджера:\n\n"
        "👩 Имя: Лола\n"
        "📱 Номер телефона: +79654138034"
    )


@dp.message(F.text == "✏️ Редактировать данные")
async def edit_data(message: Message):
    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer(
            "Сначала необходимо пройти регистрацию через команду /start.",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    await message.answer(
        "Ваши текущие данные:\n\n"
        f"ФИО: {staff['full_name']}\n"
        f"Гражданство: {staff['citizenship']}\n"
        f"Телефон: {staff['phone']}\n\n"
        "Редактирование данных мы добавим следующим этапом."
    )

@dp.message(F.text == "❌ Отказ от мероприятия")
async def refuse_event_start(message: Message):
    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer(
            "Сначала необходимо пройти регистрацию через команду /start.",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    if is_staff_blacklisted(staff["staff_id"]):
        await message.answer(
            "❌ Вы не можете отказаться от мероприятия, так как находитесь в черном списке."
        )
        return

    rows = get_staff_active_responses_with_ids(staff["staff_id"])

    if not rows:
        await message.answer("У вас нет активных записей на мероприятия.")
        return

    text = "Выберите смену, от которой хотите отказаться:\n\n"
    keyboard = []

    for row in rows:
        text += (
            f"📅 Дата: {row['event_date']}\n"
            f"🕒 Смена: {row['shift_start']}–{row['shift_end']}\n"
            f"Статус: {row['response_status']}\n"
            f"Комментарий: {row['comment']}\n\n"
        )

        keyboard.append([
            InlineKeyboardButton(
                text=f"Отказаться {row['event_date']} {row['shift_start']}",
                callback_data=f"refuse_shift:{row['shift_id']}"
            )
        ])

    await message.answer(
        text,
        reply_markup=InlineKeyboardMarkup(inline_keyboard=keyboard)
    )

@dp.callback_query(F.data.startswith("refuse_shift:"))
async def refuse_shift_callback(callback: CallbackQuery):
    staff = get_staff_by_telegram_id(callback.from_user.id)

    if not staff:
        await callback.message.answer("Сначала необходимо пройти регистрацию через команду /start.")
        await callback.answer()
        return

    shift_id = int(callback.data.split(":")[1])

    shift = get_shift_by_id(shift_id)

    if not shift:
        await callback.message.answer("Смена не найдена или уже недоступна.")
        await callback.answer()
        return

    removed = remove_staff_from_shift_response(
        shift_id=shift_id,
        staff_id=staff["staff_id"]
    )

    if not removed:
        await callback.message.answer("Вы не были записаны на эту смену.")
        await callback.answer()
        return

    await callback.message.answer(
        "✅ Вы отказались от смены.\n\n"
        f"Дата: {shift['event_date']}\n"
        f"Смена: {shift['shift_start']}–{shift['shift_end']}"
    )

    await bot.send_message(
        ADMIN_ID,
        "Сотрудник отказался от смены.\n\n"
        f"ФИО: {staff['full_name']}\n"
        f"Телефон: {staff['phone']}\n"
        f"Дата: {shift['event_date']}\n"
        f"Смена: {shift['shift_start']}–{shift['shift_end']}"
    )

    await callback.answer("Отказ оформлен")

async def auto_cleanup():
    while True:
        try:
            archive_past_events()
            clear_old_events()
        except Exception as error:
            print("Ошибка автоочистки:", error)

        await asyncio.sleep(3600)

@dp.message(F.text == "♻️ Удалить из черного списка")
async def remove_blacklist_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    blacklist = get_blacklist()

    if not blacklist:
        await message.answer("Черный список пуст.")
        return

    text = "Введите фамилию пользователя, которого нужно удалить из черного списка.\n\nТекущий черный список:\n\n"

    for i, item in enumerate(blacklist, start=1):
        username = item["telegram_username"] or "не указан"
        text += (
            f"{i}. {item['full_name']}\n"
            f"Телефон: {item['phone']}\n"
            f"Telegram: @{username}\n"
            f"Причина: {item['reason']}\n\n"
        )

    await message.answer(text)
    await state.set_state(RemoveBlacklist.lastname)


@dp.message(RemoveBlacklist.lastname)
async def remove_blacklist_by_lastname(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    staff = find_staff_by_lastname(message.text.strip())

    if not staff:
        await message.answer("Пользователь с такой фамилией не найден. Попробуйте снова.")
        return

    if not is_staff_blacklisted(staff["staff_id"]):
        await message.answer("Этот пользователь не находится в черном списке.")
        await state.clear()
        return

    remove_staff_from_blacklist(staff["staff_id"])

    try:
        await bot.send_message(
            staff["telegram_id"],
            "✅ Вы удалены из черного списка.\n\n"
            "Теперь вы снова можете записываться на мероприятия."
        )
    except Exception:
        pass

    await state.clear()

    await message.answer(
        f"Пользователь {staff['full_name']} удалён из черного списка.\n"
        "Его регистрационные данные сохранены."
    )

@dp.message(F.text == "➕ Добавить смену")
async def add_shift_start(message: Message, state: FSMContext):
    if message.from_user.id != ADMIN_ID:
        return

    events = get_active_events()

    if not events:
        await message.answer(
            "Нет активных мероприятий.\n"
            "Сначала создайте мероприятие."
        )
        return

    keyboard = []

    text = "Выберите мероприятие для добавления смены:\n\n"

    for event in events:
        text += (
            f"📅 {event['event_date']}\n"
            f"Комментарий: {event['comment']}\n\n"
        )

        keyboard.append([
            InlineKeyboardButton(
                text=f"{event['event_date']}",
                callback_data=f"add_shift_event:{event['event_id']}"
            )
        ])

    await message.answer(
        text,
        reply_markup=InlineKeyboardMarkup(inline_keyboard=keyboard)
    )

@dp.callback_query(F.data.startswith("add_shift_event:"))
async def add_shift_event_callback(callback: CallbackQuery, state: FSMContext):
    if callback.from_user.id != ADMIN_ID:
        return

    event_id = int(callback.data.split(":")[1])

    events = get_active_events()

    selected_event = None

    for event in events:
        if event["event_id"] == event_id:
            selected_event = event
            break

    if not selected_event:
        await callback.message.answer("Мероприятие не найдено.")
        await callback.answer()
        return

    await state.update_data(
        event_id=selected_event["event_id"],
        event_date=selected_event["event_date"],
        event_comment=selected_event["comment"]
    )

    await callback.message.answer(
        f"Вы выбрали мероприятие:\n\n"
        f"📅 Дата: {selected_event['event_date']}\n"
        f"Комментарий: {selected_event['comment']}\n\n"
        "Введите время начала смены.\n"
        "Например: 07:00"
    )

    await state.set_state(AddShift.shift_start)

    await callback.answer()

@dp.message(AddShift.shift_start)
async def add_shift_start_time(message: Message, state: FSMContext):
    shift_start = message.text.strip()

    if len(shift_start) == 4:
        shift_start = "0" + shift_start

    await state.update_data(shift_start=shift_start)

    await message.answer(
        "Введите время окончания смены.\n"
        "Формат: ЧЧ:ММ\n"
        "Например: 19:00"
    )
    await state.set_state(AddShift.shift_end)


@dp.message(AddShift.shift_end)
async def add_shift_end_time(message: Message, state: FSMContext):
    shift_end = message.text.strip()

    if len(shift_end) == 4:
        shift_end = "0" + shift_end

    await state.update_data(shift_end=shift_end)

    await message.answer(
        "Введите количество сотрудников на эту смену.\n"
        "Например: 10"
    )
    await state.set_state(AddShift.staff_required)


@dp.message(AddShift.staff_required)
async def add_shift_staff_required(message: Message, state: FSMContext):
    if not message.text.isdigit():
        await message.answer("Введите число. Например: 10")
        return

    data = await state.get_data()

    shift_id = create_event_shift(
        event_id=data["event_id"],
        shift_start=data["shift_start"],
        shift_end=data["shift_end"],
        staff_required=int(message.text)
    )

    staff_list = get_all_staff()

    mailing_text = (
        "📢 Добрый день, коллеги!\n\n"
        "Открыт набор персонала на смену.\n\n"
        f"Дата мероприятия: {data['event_date']}\n"
        f"Время смены: {data['shift_start']}–{data['shift_end']}\n"
        f"Нужно сотрудников: {message.text}\n\n"
        f"Комментарий: {data['event_comment']}\n\n"
        "Чтобы записаться:\n"
        "1. Нажмите кнопку «Записаться на смену»\n"
        "2. Отправьте дату и время начала смены:\n"
        f"{data['event_date']} {data['shift_start']}"
    )
    
    sent_count = 0

    for staff in staff_list:
        try:
            await bot.send_message(staff["telegram_id"], mailing_text)
            sent_count += 1
        except Exception:
            pass

    await state.clear()

    await message.answer(
        "Смена добавлена.\n\n"
        f"ID смены: {shift_id}\n"
        f"Дата: {data['event_date']}\n"
        f"Время: {data['shift_start']}–{data['shift_end']}\n"
        f"Количество сотрудников: {message.text}\n"
        f"Рассылка отправлена пользователям: {sent_count}"
    )

@dp.message(SignUpShift.shift_data)
async def shift_response_handler(message: Message, state: FSMContext):
    print("Пользователь ввёл данные для записи:", message.text)

    staff = get_staff_by_telegram_id(message.from_user.id)

    if not staff:
        await message.answer("Сначала необходимо пройти регистрацию через команду /start.")
        await state.clear()
        return

    text = message.text.strip()

    if not re.match(r"^\d{4}-\d{2}-\d{2}\s+\d{1,2}:\d{2}$", text):
        await message.answer(
            "Введите дату и время в правильном формате:\n\n"
            "2026-05-28 7:00"
        )
        return

    event_date, shift_start = text.split()

    if len(shift_start) == 4:
        shift_start = "0" + shift_start

    print("Дата:", event_date)
    print("Время:", shift_start)

    shift = get_shift_by_date_and_start(event_date, shift_start)

    print("Найденная смена:", shift)

    if not shift:
        await message.answer(
            "Смена на указанную дату и время не найдена.\n\n"
            "Проверьте данные и попробуйте снова."
        )
        return

    if has_staff_response_on_event_date(event_date, staff["staff_id"]):
        await message.answer(
            "Вы уже записаны на смену в этот день.\n\n"
            "В один день можно записаться только на одну смену."
        )
        await state.clear()
        return

    current_count = count_shift_responses(shift["shift_id"])
    required_count = shift["staff_required"]

    if current_count >= required_count:
        await message.answer("На эту смену набор уже закрыт.")
        close_shift(shift["shift_id"])
        await state.clear()
        return

    added = add_shift_response(
        event_id=shift["event_id"],
        shift_id=shift["shift_id"],
        staff_id=staff["staff_id"]
    )

    print("Результат добавления:", added)

    if not added:
        await message.answer("Вы уже записались на эту смену.")
        await state.clear()
        return

    current_count = count_shift_responses(shift["shift_id"])

    await state.clear()

    await message.answer(
        "✅ Вы записались на смену.\n\n"
        f"Дата мероприятия: {shift['event_date']}\n"
        f"Время смены: {shift['shift_start']}–{shift['shift_end']}\n"
        f"Сейчас записалось: {current_count} из {required_count}",
        reply_markup=user_keyboard()
    )

    await bot.send_message(
        ADMIN_ID,
        "Новый отклик на смену.\n\n"
        f"ФИО: {staff['full_name']}\n"
        f"Телефон: {staff['phone']}\n"
        f"Дата: {shift['event_date']}\n"
        f"Смена: {shift['shift_start']}–{shift['shift_end']}\n"
        f"Записалось: {current_count} из {required_count}"
    )

    if current_count >= required_count:
        close_shift(shift["shift_id"])

        staff_list = get_all_staff()

        for person in staff_list:
            try:
                await bot.send_message(
                    person["telegram_id"],
                    "🔒 Набор на смену закрыт.\n\n"
                    f"Дата мероприятия: {shift['event_date']}\n"
                    f"Смена: {shift['shift_start']}–{shift['shift_end']}\n\n"
                    "Свободных мест больше нет."
                )
            except Exception as e:
                print("Ошибка рассылки:", e)

        await bot.send_message(
            ADMIN_ID,
            "🔒 Набор персонала завершён.\n\n"
            f"Дата: {shift['event_date']}\n"
            f"Смена: {shift['shift_start']}–{shift['shift_end']}\n\n"
            "Уведомление отправлено всем пользователям."
        )

@dp.message(F.text == "📋 Все мероприятия")
async def all_events_handler(message: Message):
    if message.from_user.id != ADMIN_ID:
        return

    rows = get_all_events_with_shifts()

    if not rows:
        await message.answer("Мероприятий пока нет.")
        return

    text = "Все мероприятия и смены:\n\n"

    current_event_id = None

    for row in rows:
        if row["event_id"] != current_event_id:
            current_event_id = row["event_id"]

            status = "активно" if row["event_active"] else "закрыто"

            text += (
                f"Дата: {row['event_date']}\n"
                f"Статус мероприятия: {status}\n"
                f"Комментарий: {row['comment']}\n"
            )

        if row["shift_id"]:
            shift_status = "активна" if row["shift_active"] else "закрыта"

            text += (
                f"  Смена: {row['shift_start']}–{row['shift_end']}\n"
                f"  Нужно: {row['staff_required']}\n"
                f"  Записалось: {row['response_count']}\n"
                f"  Статус смены: {shift_status}\n\n"
            )
        else:
            text += "  Смены пока не добавлены.\n\n"

    await message.answer(text)

def create_staff_docx(event_date: str, rows):
    file_name = f"staff_list_{event_date}.docx"

    document = Document()

    document.add_heading("Список персонала мероприятия", level=1)
    document.add_paragraph(f"Дата мероприятия: {event_date}")

    for row in rows:
        if not row["shift_id"]:
            continue

        participants = get_shift_participants(row["shift_id"])

        document.add_heading(
            f"Смена {row['shift_start']}–{row['shift_end']}",
            level=2
        )

        document.add_paragraph(f"Нужно сотрудников: {row['staff_required']}")
        document.add_paragraph(f"Записалось: {len(participants)}")

        if not participants:
            document.add_paragraph("На смену пока никто не записался.")
            continue

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        header = table.rows[0].cells
        header[0].text = "№"
        header[1].text = "ФИО"
        header[2].text = "Телефон"
        header[3].text = "Статус"

        for i, participant in enumerate(participants, start=1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = participant["full_name"]
            cells[2].text = participant["phone"]
            cells[3].text = participant["response_status"]

    document.save(file_name)
    return file_name

def create_staff_xlsx(event_date: str, rows):
    file_name = f"staff_list_{event_date}.xlsx"

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Список персонала"

    sheet["A1"] = "Список персонала мероприятия"
    sheet["A1"].font = Font(bold=True, size=16)

    sheet["A2"] = f"Дата мероприятия: {event_date}"

    row_number = 4

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    for row in rows:
        if not row["shift_id"]:
            continue

        participants = get_shift_participants(row["shift_id"])

        sheet.cell(row=row_number, column=1).value = f"Смена {row['shift_start']}–{row['shift_end']}"
        sheet.cell(row=row_number, column=1).font = Font(bold=True, size=13)
        row_number += 1

        sheet.cell(row=row_number, column=1).value = "Нужно сотрудников"
        sheet.cell(row=row_number, column=2).value = row["staff_required"]
        row_number += 1

        sheet.cell(row=row_number, column=1).value = "Записалось"
        sheet.cell(row=row_number, column=2).value = len(participants)
        row_number += 2

        headers = ["№", "ФИО", "Телефон", "Статус"]

        for col, header in enumerate(headers, start=1):
            cell = sheet.cell(row=row_number, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border

        row_number += 1

        if participants:
            for i, participant in enumerate(participants, start=1):
                values = [
                    i,
                    participant["full_name"],
                    participant["phone"],
                    participant["response_status"]
                ]

                for col, value in enumerate(values, start=1):
                    cell = sheet.cell(row=row_number, column=col)
                    cell.value = value
                    cell.border = thin_border

                row_number += 1
        else:
            sheet.cell(row=row_number, column=1).value = "На смену пока никто не записался."
            row_number += 1

        row_number += 2

    sheet.column_dimensions["A"].width = 8
    sheet.column_dimensions["B"].width = 35
    sheet.column_dimensions["C"].width = 20
    sheet.column_dimensions["D"].width = 18

    workbook.save(file_name)
    return file_name

async def auto_shift_reminders():
    while True:
        try:
            reminders = get_tomorrow_shift_participants_for_reminder()

            for item in reminders:
                try:
                    await bot.send_message(
                        item["telegram_id"],
                        "⏰ Напоминание о смене!\n\n"
                        "Завтра у вас запланирована смена.\n\n"
                        f"Дата: {item['event_date']}\n"
                        f"Смена: {item['shift_start']}–{item['shift_end']}\n"
                        f"Комментарий: {item['comment']}\n\n"
                        "Пожалуйста, не опаздывайте."
                    )

                    mark_shift_reminder_sent(item["response_id"], "tomorrow")

                except Exception:
                    pass

        except Exception as error:
            print("Ошибка при отправке напоминаний:", error)

        await asyncio.sleep(3600)

async def main():
    print("Бот запущен...")
    asyncio.create_task(auto_cleanup())
    asyncio.create_task(auto_shift_reminders())
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
